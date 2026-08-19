import os
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("/Volumes/mark/rwkvrag/Windows本地文档智能搜索_实现方案与技术建议_20260814.docx")
DOC_FONT = os.environ.get("DOC_FONT", "Microsoft YaHei")

BLUE = "2E74B5"
NAVY = "17365D"
LIGHT_BLUE = "EAF2F8"
PALE_BLUE = "F4F8FC"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "D7DEE7"
DARK_GRAY = "4A5560"
GREEN = "2E7D32"
PALE_GREEN = "EAF4EA"
AMBER = "9A6700"
PALE_AMBER = "FFF4CE"
WHITE = "FFFFFF"
BLACK = "1F1F1F"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=MID_GRAY, size=6):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    total = sum(widths_dxa)
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths_dxa[idx] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tr_pr.append(cant_split)


def set_run_font(run, size=None, bold=None, color=None, italic=None, latin=None, east=None):
    latin = latin or DOC_FONT
    east = east or DOC_FONT
    run.font.name = latin
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), east)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def format_all_runs(paragraph, size=10.5, color=BLACK, bold=None):
    for run in paragraph.runs:
        set_run_font(run, size=size, color=color, bold=bold)


def add_body(doc, text, bold_lead=None, after=6):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        set_run_font(r1, size=10.5, bold=True, color=NAVY)
        r2 = p.add_run(text[len(bold_lead):])
        set_run_font(r2, size=10.5, color=BLACK)
    else:
        r = p.add_run(text)
        set_run_font(r, size=10.5, color=BLACK)
    return p


def add_bullet(doc, text, level=0, bold_lead=None):
    style = "List Bullet" if level == 0 else "List Bullet 2"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.12
    if bold_lead and text.startswith(bold_lead):
        a = p.add_run(bold_lead)
        set_run_font(a, size=10.5, bold=True, color=NAVY)
        b = p.add_run(text[len(bold_lead):])
        set_run_font(b, size=10.5, color=BLACK)
    else:
        r = p.add_run(text)
        set_run_font(r, size=10.5, color=BLACK)
    return p


def new_numbering(doc):
    numbering = doc.part.numbering_part.element
    nums = numbering.findall(qn("w:num"))
    next_id = max((int(n.get(qn("w:numId"))) for n in nums), default=0) + 1

    style_num_id = doc.styles["List Number"]._element.pPr.numPr.numId.val
    base_num = next(n for n in nums if int(n.get(qn("w:numId"))) == style_num_id)
    abstract_id = base_num.find(qn("w:abstractNumId")).get(qn("w:val"))

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(next_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), abstract_id)
    num.append(abstract)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    num.append(override)
    numbering.append(num)
    return next_id


def add_number(doc, text, num_id):
    p = doc.add_paragraph(style="List Number")
    num_pr = p._p.get_or_add_pPr().get_or_add_numPr()
    num_pr.get_or_add_ilvl().val = 0
    num_pr.get_or_add_numId().val = num_id
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.12
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=BLACK)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    return p


def add_callout(doc, label, text, fill=LIGHT_BLUE, accent=BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [9360])
    set_table_borders(table, color=accent, size=7)
    set_row_cant_split(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.08
    r1 = p.add_run(label + "  ")
    set_run_font(r1, size=10.5, bold=True, color=accent)
    r2 = p.add_run(text)
    set_run_font(r2, size=10.5, color=BLACK)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)
    return table


def add_table(doc, headers, rows, widths, header_fill=LIGHT_GRAY, font_size=9.3, aligns=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, widths)
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        set_run_font(r, size=font_size, bold=True, color=NAVY)
    for row_data in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_data):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.06
            if aligns:
                p.alignment = aligns[i]
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(value))
            set_run_font(r, size=font_size, color=BLACK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, text, fld_end])
    set_run_font(run, size=8.5, color=DARK_GRAY)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.78)
section.bottom_margin = Inches(0.72)
section.left_margin = Inches(1.0)
section.right_margin = Inches(1.0)
section.header_distance = Inches(0.32)
section.footer_distance = Inches(0.32)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = DOC_FONT
normal._element.rPr.rFonts.set(qn("w:ascii"), DOC_FONT)
normal._element.rPr.rFonts.set(qn("w:hAnsi"), DOC_FONT)
normal._element.rPr.rFonts.set(qn("w:eastAsia"), DOC_FONT)
normal.font.size = Pt(10.5)
normal.font.color.rgb = RGBColor.from_string(BLACK)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.10

for name, size, color, before, after in (
    ("Heading 1", 16, BLUE, 16, 8),
    ("Heading 2", 13, BLUE, 12, 6),
    ("Heading 3", 11.5, NAVY, 8, 4),
):
    st = styles[name]
    st.font.name = DOC_FONT
    st._element.rPr.rFonts.set(qn("w:ascii"), DOC_FONT)
    st._element.rPr.rFonts.set(qn("w:hAnsi"), DOC_FONT)
    st._element.rPr.rFonts.set(qn("w:eastAsia"), DOC_FONT)
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(color)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

for style_name in ("List Bullet", "List Bullet 2", "List Number"):
    st = styles[style_name]
    st.font.name = DOC_FONT
    st._element.rPr.rFonts.set(qn("w:ascii"), DOC_FONT)
    st._element.rPr.rFonts.set(qn("w:hAnsi"), DOC_FONT)
    st._element.rPr.rFonts.set(qn("w:eastAsia"), DOC_FONT)
    st.font.size = Pt(10.5)

header = section.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
hr = hp.add_run("技术方案简报  |  Windows 本地文档智能搜索")
set_run_font(hr, size=8.5, color=DARK_GRAY)

footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = fp.add_run("内部评估与供应商沟通参考  |  ")
set_run_font(fr, size=8.5, color=DARK_GRAY)
add_page_field(fp)

# First-page masthead
sp = doc.add_paragraph()
sp.paragraph_format.space_after = Pt(14)

kicker = doc.add_paragraph()
kicker.paragraph_format.space_after = Pt(3)
kr = kicker.add_run("IMPLEMENTATION BRIEF / 实现方案")
set_run_font(kr, size=9.5, bold=True, color=BLUE)

title = doc.add_paragraph()
title.paragraph_format.space_after = Pt(5)
tr = title.add_run("Windows 本地文档智能搜索")
set_run_font(tr, size=25, bold=True, color=NAVY)

subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_after = Pt(15)
sr = subtitle.add_run("实现方案与技术建议")
set_run_font(sr, size=15, bold=True, color=DARK_GRAY)

meta = [
    ("依据", "《Windows 本地文档智能搜索需求文档（供应商评估版）》"),
    ("方案定位", "Windows 10/11，本地优先，支持 5,000-50,000 个文件"),
    ("推荐版本", "MVP：混合检索；扩展：OCR、Reranker、本地问答"),
    ("日期", "2026-08-14"),
]
for label, value in meta:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    a = p.add_run(label + "：")
    set_run_font(a, size=10.5, bold=True, color=NAVY)
    b = p.add_run(value)
    set_run_font(b, size=10.5, color=BLACK)

rule = doc.add_paragraph()
rule.paragraph_format.space_before = Pt(8)
rule.paragraph_format.space_after = Pt(14)
ppr = rule._p.get_or_add_pPr()
pbdr = OxmlElement("w:pBdr")
bottom = OxmlElement("w:bottom")
bottom.set(qn("w:val"), "single")
bottom.set(qn("w:sz"), "18")
bottom.set(qn("w:space"), "1")
bottom.set(qn("w:color"), BLUE)
pbdr.append(bottom)
ppr.append(pbdr)

add_callout(
    doc,
    "核心结论",
    "本项目应优先建设“本地文档解析 + 关键词检索 + Embedding 向量检索 + 融合排序”的混合搜索系统。MVP 不依赖生成式大模型；只有启用问答摘要时才需要本地 LLM。",
    fill=PALE_GREEN,
    accent=GREEN,
)

add_heading(doc, "建议优先确认的 5 项决策", 1)
for item in (
    "默认完全本地运行，文档正文、查询和索引不上传外部云端。",
    "MVP 使用混合检索，不能只做向量搜索，也不能只做 Windows 文件名搜索。",
    "推荐技术基线：.NET 8/WPF + 本地后台服务 + ONNX Runtime + SQLite FTS5 + HNSW。",
    "OCR、生成式问答、企业权限同步和多人共享服务不纳入首期必选范围。",
    "搜索质量以客户真实查询集验收，至少评测 Recall@10、MRR@10 和 Top 10 命中率。",
):
    add_bullet(doc, item)

add_heading(doc, "预期成果", 1)
add_body(doc, "形成可安装、可维护、可离线运行的 Windows 工具：普通用户可选择目录、查看索引进度、使用自然语言或关键词搜索，并从结果直接追溯和打开原文件。")

add_page_break(doc)

add_heading(doc, "1. 需求分析与范围界定", 1)
add_heading(doc, "1.1 核心业务目标", 2)
add_body(doc, "项目要解决的不是“把文件交给大模型聊天”，而是在本地办公文档中稳定地找到相关文件、相关段落和来源位置。成功条件由四部分共同决定：解析覆盖率、索引更新可靠性、搜索相关性和结果可追溯性。")

add_table(
    doc,
    ["能力层", "MVP 必须具备", "扩展能力"],
    [
        ("数据接入", "本地目录、桌面、文档目录、映射盘", "企业共享盘、集中式文件库"),
        ("格式处理", "PDF、DOC/DOCX、XLS/XLSX、PPT/PPTX、TXT、MD", "图片、邮件、压缩包、特殊格式"),
        ("检索", "关键词、语义、过滤、融合排序", "Reranker、跨语言增强、查询改写"),
        ("结果", "文件、路径、片段、相关度、时间、打开入口", "页码/工作表/幻灯片精确定位"),
        ("智能化", "不依赖生成式 LLM", "OCR、问答摘要、本地 RAG"),
    ],
    [1700, 4100, 3560],
)

add_heading(doc, "1.2 关键约束", 2)
for text, lead in (
    ("本地优先：默认不上传文档正文、查询和索引；任何联网能力都必须显式授权。", "本地优先："),
    ("办公电脑适配：应在无独立 GPU 的常见企业电脑上可用，并允许控制后台 CPU 和内存占用。", "办公电脑适配："),
    ("规模口径：5 万“文件”并不足以定义负载，还要统计总容量、总页数、片段数和扫描件比例。", "规模口径："),
    ("三语场景：界面要求英文、俄文、中文；需进一步确认俄文是否也属于内容语义检索范围。", "三语场景："),
):
    add_bullet(doc, text, bold_lead=lead)

add_callout(doc, "范围控制", "建议首期不纳入 OCR、生成式问答、企业级权限同步、多人服务器部署和复杂公式/图片理解。这些能力会显著增加周期、资源占用与验收复杂度。", fill=PALE_AMBER, accent=AMBER)

add_heading(doc, "1.3 对现有需求文档的判断", 2)
add_body(doc, "现有文档足以支持供应商初步评估和报价，但仍属于需求边界说明，不宜直接作为最终开发规格。第 4 节后直接进入第 6 节，需确认第 5 节是否遗漏；性能、安全、权限和检索质量指标也需在详细设计阶段量化。")

add_page_break(doc)

add_heading(doc, "2. 推荐总体架构", 1)
add_callout(doc, "架构原则", "桌面端负责交互，本地后台负责解析、推理和索引；所有核心数据存放在本机，服务只监听 127.0.0.1。", fill=LIGHT_BLUE, accent=BLUE)

add_heading(doc, "2.1 逻辑架构", 2)
arch = add_table(
    doc,
    ["Windows 桌面端", "本地后台服务", "本地存储"],
    [
        ("目录配置\n自然语言/关键词查询\n结果预览与打开\n索引状态和失败列表", "文件扫描与变更监听\n格式解析与文本清洗\n分段与 Embedding\n混合召回与排序\n任务队列和错误重试", "SQLite 元数据\nFTS5 关键词索引\nHNSW 向量索引\n模型文件、配置和日志"),
    ],
    [2800, 3560, 3000],
    header_fill=LIGHT_BLUE,
    font_size=9.6,
    aligns=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
)

add_heading(doc, "2.2 核心数据流", 2)
flow_num = new_numbering(doc)
for step in (
    "扫描指定目录，建立文件清单并读取路径、类型、大小、修改时间和权限状态。",
    "按格式解析正文、标题、表格和位置元数据，失败文件进入可重试队列。",
    "按文档结构分段，生成关键词索引项和 Embedding 向量。",
    "查询时并行执行关键词召回与向量召回，再通过 RRF 或加权规则融合。",
    "返回文件级和片段级结果，提供复制路径、打开文件和打开所在目录。",
):
    add_number(doc, step, flow_num)

add_heading(doc, "2.3 部署形态", 2)
add_table(
    doc,
    ["组件", "推荐部署方式", "关键要求"],
    [
        ("桌面客户端", "WPF 桌面应用", "单实例、托盘运行、多语言、自动恢复"),
        ("后台服务", "随客户端安装的本地进程/Windows 服务", "仅本机访问、低优先级后台任务"),
        ("模型", "安装包内置或离线模型包", "版本固定、哈希校验、可替换升级"),
        ("索引", "用户目录或管理员指定目录", "可迁移、可清理、版本可识别"),
    ],
    [1900, 3200, 4260],
)

add_page_break(doc)

add_heading(doc, "3. 技术选型建议", 1)
add_table(
    doc,
    ["领域", "首选技术", "选择理由 / 注意事项"],
    [
        ("桌面端", ".NET 8 + WPF + MVVM", "Windows 集成成熟；适合企业安装、托盘、Shell 操作和本地服务管理"),
        ("安装包", "MSIX 或 WiX Toolset", "支持签名、升级和企业分发；需兼容无管理员权限场景"),
        ("后台", ".NET Worker 或 Python 3.11/3.12", "若团队偏 C#，优先全 .NET；若解析生态优先，可打包独立 Python 运行时"),
        ("本地 API", "Named Pipe 或 FastAPI/HTTP", "HTTP 仅绑定 127.0.0.1；应增加随机令牌或进程级访问控制"),
        ("Embedding", "bge-m3 + ONNX Runtime", "中文、英文、俄文及混合内容适配较好；CPU 可运行"),
        ("关键词索引", "SQLite FTS5", "MVP 部署简单；适合本机几十万至约百万片段级索引"),
        ("向量索引", "HNSWlib", "单机检索速度快、依赖较轻；模型升级时支持重建"),
        ("元数据", "SQLite", "事务可靠、易迁移和备份；保存文件、片段、任务和错误状态"),
        ("OCR（可选）", "PaddleOCR", "用于扫描 PDF 和图片；应独立计费和配置资源上限"),
        ("问答（可选）", "本地 LLM + RAG", "只基于检索结果生成，必须引用来源；不属于语义搜索的前置条件"),
    ],
    [1550, 2650, 5160],
    font_size=8.9,
)

add_heading(doc, "3.1 Embedding 模型策略", 2)
add_body(doc, "推荐将 bge-m3 作为首轮技术验证模型，但最终选型必须基于客户真实语料测试，而不是只看公开榜单。模型文件、向量维度、归一化规则和分词方式应作为索引版本的一部分；更换模型后必须重建向量索引。")
for item in (
    "默认使用 ONNX Runtime CPU 推理，批量索引时控制线程数，降低对办公应用的干扰。",
    "检测到受支持 GPU 后可启用 DirectML 或 CUDA，但不能把 GPU 设为运行前提。",
    "可对索引向量使用 FP16 或量化降低存储，但应通过相关性评测确认精度损失。",
):
    add_bullet(doc, item)

add_callout(doc, "为什么不只用 Embedding", "文件名、编号、日期、合同条款号和精确短语更适合关键词检索；同义表达和自然语言描述更适合向量检索。两者融合才能覆盖真实办公搜索。", fill=PALE_GREEN, accent=GREEN)

add_page_break(doc)

add_heading(doc, "4. 文档解析、分段与索引", 1)
add_heading(doc, "4.1 格式解析", 2)
add_table(
    doc,
    ["格式", "推荐实现", "重点保留信息"],
    [
        ("DOCX", "Open XML SDK / python-docx", "标题层级、正文、表格、段落顺序"),
        ("DOC", "LibreOffice 转换或商业组件", "旧格式兼容性、转换失败日志"),
        ("XLSX", "Open XML SDK / openpyxl", "工作表、连续区域、公式显示值、单元格位置"),
        ("XLS", "LibreOffice 或兼容库", "编码、日期、合并单元格"),
        ("PPTX", "Open XML SDK / python-pptx", "幻灯片号、标题、文本框和备注"),
        ("PPT", "LibreOffice 转换", "旧格式兼容性"),
        ("PDF", "PyMuPDF", "页码、文本块、阅读顺序、加密状态"),
        ("TXT/MD", "编码检测 + 原生读取", "标题、段落、代码块、文件编码"),
        ("扫描 PDF", "PaddleOCR（可选）", "页码、识别置信度、语言模型"),
    ],
    [1300, 3000, 5060],
    font_size=9.0,
)

add_heading(doc, "4.2 分段策略", 2)
add_body(doc, "不能对所有格式统一按固定字符硬切。应优先尊重标题、段落、页面、工作表和幻灯片结构，在块过长时再执行二次切分。")
for text, lead in (
    ("Word：按标题路径、段落和表格组织；标题路径作为每个片段的上下文。", "Word："),
    ("PDF：按页和版面文本块组织，保留页码；处理页眉页脚重复文本。", "PDF："),
    ("Excel：按工作表和连续表格区域组织，避免把整张表压成一个超长片段。", "Excel："),
    ("PPT：按幻灯片组织，标题和正文合并，保留幻灯片号。", "PPT："),
    ("通用参数：建议每片段 300-700 个中文字符，重叠 50-100 字，再通过评测调整。", "通用参数："),
):
    add_bullet(doc, text, bold_lead=lead)

add_heading(doc, "4.3 片段元数据", 2)
add_body(doc, "每个片段至少保存：chunk_id、file_id、绝对路径、文件名、类型、修改时间、大小、内容哈希、页码/工作表/幻灯片号、标题路径、片段序号、片段文本、解析器版本和 Embedding 模型版本。")

add_page_break(doc)
add_heading(doc, "4.4 旧 Office 格式的决策", 2)
add_callout(doc, "需提前确认", "DOC/XLS/PPT 旧格式若要求高可靠且不允许依赖本机 Office，建议采购商业解析组件；采用 LibreOffice 可降低许可成本，但应在目标企业环境验证安装体积、杀毒软件兼容性和转换稳定性。", fill=PALE_AMBER, accent=AMBER)

add_heading(doc, "5. 搜索与排序实现", 1)
add_heading(doc, "5.1 查询流程", 2)
flow_rows = [
    ("1. 查询解析", "识别自然语言、精确短语、文件类型、时间范围和路径过滤"),
    ("2. 并行召回", "FTS5 关键词召回 + 查询 Embedding/HNSW 向量召回"),
    ("3. 结果融合", "使用 RRF 或经验证的加权规则合并排名"),
    ("4. 可选重排", "对前 30-100 个候选使用 bge-reranker-v2-m3 重排"),
    ("5. 文件聚合", "同一文件限制片段数量，兼顾最相关片段与结果多样性"),
    ("6. 结果展示", "文件名、路径、片段、位置、时间、类型、相关度和打开操作"),
]
add_table(doc, ["阶段", "实现要点"], flow_rows, [2100, 7260], header_fill=LIGHT_BLUE, font_size=9.5)

add_heading(doc, "5.2 排序建议", 2)
add_body(doc, "MVP 推荐使用 RRF 融合，因为它对不同召回分数尺度不敏感，调参成本较低。文件名精确命中、路径命中和时间过滤应作为显式规则处理，不要全部交给向量相似度。")
for item in (
    "精确文件名、完整短语或编号命中应获得明显加权。",
    "语义召回负责覆盖同义词、自然语言描述和跨段落表达差异。",
    "相关度展示不宜直接暴露原始余弦分数，可显示高/中/低或经过校准的百分比。",
    "Reranker 是提升 Top 10 排序质量的增强项，但会增加 CPU 延迟和模型体积。",
):
    add_bullet(doc, item)

add_heading(doc, "5.3 搜索质量评测", 2)
add_callout(doc, "验收底线", "没有真实查询集和人工相关性标注，“能返回语义相关结果”只能主观判断，无法形成可执行的供应商验收标准。", fill=PALE_GREEN, accent=GREEN)
add_table(
    doc,
    ["指标", "用途", "建议验收方式"],
    [
        ("Recall@10", "前 10 条是否覆盖相关材料", "以人工标注相关文件/片段计算"),
        ("MRR@10", "首个正确结果是否靠前", "重点衡量常用查询体验"),
        ("nDCG@10", "多级相关性排序质量", "对高度相关/部分相关分级标注"),
        ("Top 10 命中率", "业务可理解的成功率", "至少 100 条真实查询统计"),
    ],
    [1500, 3150, 4710],
)

add_heading(doc, "6. 增量索引与任务调度", 1)
add_heading(doc, "6.1 更新机制", 2)
update_num = new_numbering(doc)
for step in (
    "首次扫描生成文件清单，记录路径、大小、修改时间和快速指纹。",
    "使用 FileSystemWatcher 监听新增、修改、删除和重命名事件。",
    "对短时间内的重复事件进行 2-5 秒去抖和合并。",
    "通过周期性全量校验补偿文件监听可能丢失的事件。",
    "修改文件时删除旧片段后重新解析和向量化；删除文件时同步清理索引。",
    "模型或解析器版本变化时，将受影响数据标记为待重建。",
):
    add_number(doc, step, update_num)

add_heading(doc, "6.2 任务状态模型", 2)
add_table(
    doc,
    ["状态", "含义", "处理方式"],
    [
        ("Pending", "等待解析或向量化", "按优先级和资源上限调度"),
        ("Processing", "正在处理", "记录开始时间、进程和阶段"),
        ("Succeeded", "索引成功", "记录版本和完成时间"),
        ("Retryable", "被占用、临时 I/O 错误", "指数退避并限制重试次数"),
        ("Failed", "格式损坏、不支持或权限不足", "展示明确原因，允许手工重试"),
        ("Deleted", "源文件已删除", "事务性清理元数据和索引"),
    ],
    [1700, 3400, 4260],
)

add_heading(doc, "6.3 办公体验保护", 2)
for item in (
    "索引线程数、批大小和 CPU 使用上限可配置；系统繁忙时自动降速。",
    "首次索引显示文件数、片段数、成功/失败数量、预计进度和取消入口。",
    "查询优先级高于后台索引，保证用户搜索不被大批量 Embedding 阻塞。",
    "断电或应用退出后可继续任务，避免每次从头重建。",
):
    add_bullet(doc, item)

add_heading(doc, "7. 安全、隐私与权限", 1)
add_callout(doc, "安全目标", "系统不修改原始文档，不在未授权情况下上传内容；索引、缓存、日志和模型调用都纳入同一数据安全边界。", fill=PALE_GREEN, accent=GREEN)

add_table(
    doc,
    ["控制点", "推荐措施"],
    [
        ("网络", "默认无外部 API；本地服务仅监听 127.0.0.1；联网功能单独授权并可审计"),
        ("原文件", "只读访问，不覆盖、不转换回写；临时文件在任务完成后清理"),
        ("索引", "存放在用户或管理员指定目录；支持迁移、版本检查和一键彻底清理"),
        ("访问权限", "使用当前 Windows 用户权限访问本地和共享盘；权限变化后重新校验"),
        ("日志", "默认不记录正文和完整查询；路径可脱敏；日志容量和保留期可配置"),
        ("密钥", "需要加密时使用 DPAPI 保护密钥；明确设备迁移和恢复策略"),
        ("软件供应链", "安装包、模型和更新包使用数字签名及哈希校验"),
        ("遥测", "默认关闭或明确征得同意；崩溃报告不得携带正文、索引片段或查询"),
    ],
    [2100, 7260],
    font_size=9.2,
)

add_heading(doc, "7.1 关于“索引加密”", 2)
add_body(doc, "索引中可能包含正文片段、文件路径、关键词词典和向量，全部属于敏感数据。仅给 SQLite 加密而让向量文件、缓存或临时文件保持明文，不能宣称“索引已加密”。如合同要求静态加密，应覆盖所有持久化组件，并评估性能、恢复和检索时解密方式。")

add_heading(doc, "7.2 共享盘权限", 2)
add_body(doc, "映射盘和 UNC 路径应沿用当前 Windows 用户凭据。MVP 不应使用一个高权限服务账户集中抓取全部共享文件，否则本地索引可能暴露用户原本无权访问的内容。")

add_page_break(doc)

add_heading(doc, "8. 性能、容量与兼容性", 1)
add_heading(doc, "8.1 建议基线环境", 2)
add_table(
    doc,
    ["配置", "CPU / 内存", "磁盘", "适用范围"],
    [
        ("最低配置", "4 核 / 8 GB", "10 GB 可用 SSD", "小规模试用，后台索引限速"),
        ("推荐配置", "8 核 / 16 GB", "20 GB 以上 SSD", "常规 5,000-50,000 文件场景"),
    ],
    [1500, 2300, 2300, 3260],
)

add_heading(doc, "8.2 建议 MVP 指标", 2)
add_table(
    doc,
    ["指标", "建议目标", "说明"],
    [
        ("关键词查询 P95", "< 1 秒", "在定义好的基准数据集和机器上测试"),
        ("混合查询 P95", "< 3 秒", "不含可选生成式问答"),
        ("搜索首屏", "< 2 秒", "先展示主要结果，可延迟补充信息"),
        ("增量可检索时延", "普通文档 < 60 秒", "从文件稳定写入完成后开始计时"),
        ("空闲内存", "目标 < 1 GB", "模型可按需加载或常驻，需实际测量"),
        ("原文件保护", "零修改", "通过哈希或时间戳抽查验证"),
        ("未授权上传", "零外传", "通过网络抓包和日志审计验证"),
    ],
    [2200, 2300, 4860],
    font_size=9.1,
)

add_heading(doc, "8.3 容量估算", 2)
add_body(doc, "向量空间按“片段数 × 向量维度 × 每维字节数”估算。例如 50 万个 1024 维 FP32 向量，裸向量约 2 GB；再加 HNSW 图结构、正文和元数据，实际占用将达到数 GB。应在技术验证阶段采集每种文档的平均片段数。")

add_callout(doc, "不要只承诺文件数", "同样是 5 万个文件，短 TXT 与数百页扫描 PDF 的索引成本可能相差几个数量级。验收规模应同时定义文件数、总容量、总页数、片段数及扫描件比例。", fill=PALE_AMBER, accent=AMBER)

add_page_break(doc)

add_heading(doc, "9. 实施计划与交付物", 1)
add_heading(doc, "9.1 建议周期：8-12 周", 2)
add_table(
    doc,
    ["阶段", "周期", "主要产出"],
    [
        ("技术验证", "1-2 周", "样本文档解析、Embedding、索引性能与风险验证"),
        ("索引核心", "2-3 周", "扫描、解析、分段、任务队列和增量更新"),
        ("搜索核心", "2 周", "FTS、向量检索、融合排序和质量评测"),
        ("Windows 客户端", "2 周", "目录配置、搜索、结果、状态和错误页面"),
        ("联调测试", "1-2 周", "安装包、兼容性、性能、安全和缺陷修复"),
        ("试点验收", "1 周", "客户语料评测、培训、交付和遗留项确认"),
    ],
    [2100, 1500, 5760],
)

add_heading(doc, "9.2 建议交付物", 2)
for item in (
    "可安装的 Windows MVP 软件包和离线模型包。",
    "部署、升级、卸载、索引迁移和清理说明。",
    "面向普通用户的目录配置、搜索和索引管理手册。",
    "技术架构、数据流、模型版本、索引结构和安全机制说明。",
    "功能、性能、兼容性、安全及检索质量测试报告。",
    "源代码或明确的托管、许可、维护和后续升级方式。",
):
    add_bullet(doc, item)

add_heading(doc, "9.3 甲方配合事项", 2)
for item in (
    "提供脱敏后的真实文档样本和至少 100 条真实搜索查询。",
    "提供 Windows 10/11 测试机、企业杀毒软件和共享盘环境。",
    "确认旧 Office 格式、OCR、离线安装、索引加密和精确跳转要求。",
    "安排业务人员完成相关性标注并参加阶段验收。",
):
    add_bullet(doc, item)

add_page_break(doc)

add_heading(doc, "10. 验收建议", 1)
add_table(
    doc,
    ["验收领域", "检查内容", "建议证据"],
    [
        ("安装与兼容", "Windows 10/11 安装、启动、卸载、重启恢复", "兼容性测试记录"),
        ("格式覆盖", "常见文档可解析，失败有明确原因", "格式样本集与解析成功率"),
        ("搜索质量", "自然语言、关键词、文件名和过滤条件有效", "标注查询集及 Recall/MRR/nDCG"),
        ("性能", "首屏、P95、首次索引和增量时延", "同一基准机器实测报告"),
        ("增量更新", "新增、修改、删除、移动和重命名正确", "自动化用例与索引核对"),
        ("可追溯", "结果显示来源并可打开原文件/目录", "操作录屏或测试记录"),
        ("安全", "原文件零修改、未授权零上传、日志不泄露正文", "哈希比对、抓包和日志审计"),
        ("可维护", "错误列表、重试、重建、迁移和清理可用", "运维手册及演练记录"),
    ],
    [1800, 4300, 3260],
    font_size=8.9,
)

add_heading(doc, "10.1 搜索质量测试集", 2)
add_body(doc, "建议覆盖精确名称、编号、日期、同义表达、自然语言描述、中英混合、俄文内容（如纳入范围）、跨文件重复内容和无答案查询。每条查询由业务人员标注高度相关、部分相关和不相关结果。")

add_heading(doc, "10.2 无答案与低置信度", 2)
add_body(doc, "系统应允许返回“未找到足够相关内容”，而不是强行把低相似度结果包装成高相关。若启用问答摘要，答案必须附带来源文件和片段，并明确无法从检索结果确认的内容。")

add_page_break(doc)

add_heading(doc, "11. 待澄清事项与主要风险", 1)
add_table(
    doc,
    ["优先级", "待确认问题", "对方案的影响"],
    [
        ("高", "总容量、总页数、平均文件大小和扫描件比例是多少？", "决定首次索引时间、磁盘和 OCR 成本"),
        ("高", "是否允许安装 LibreOffice、模型和独立运行时？", "决定旧 Office 格式方案和安装包体积"),
        ("高", "是否必须完全离线安装和升级？", "决定模型分发、签名和更新机制"),
        ("高", "索引是否必须加密，迁移时如何恢复密钥？", "决定存储架构、性能和运维流程"),
        ("高", "搜索质量目标和真实测试集由谁提供？", "决定是否可客观验收"),
        ("中", "共享盘是映射盘、UNC、OneDrive 还是 SharePoint？", "决定监听、权限和同步方式"),
        ("中", "俄文仅为界面语言，还是也要求俄文内容检索？", "决定模型评测和测试集"),
        ("中", "是否要求 PDF 页码、Excel 工作表等精确跳转？", "决定解析元数据和客户端集成"),
        ("中", "是否允许持久化保存抽取出的正文？", "决定片段预览、加密和空间占用"),
        ("低", "原需求文档第 5 节是否遗漏？", "可能存在未表达的业务或技术范围"),
    ],
    [1100, 4920, 3340],
    font_size=8.7,
)

add_heading(doc, "11.1 主要技术风险", 2)
for text, lead in (
    ("解析差异：复杂 PDF 阅读顺序、旧 Office 格式、合并单元格和特殊编码可能导致内容缺失。", "解析差异："),
    ("质量风险：模型和分段参数若未经业务语料评测，语义搜索可能“看起来聪明但不稳定”。", "质量风险："),
    ("资源风险：大批量 Embedding 可能影响办公体验，需要限速、暂停和断点续传。", "资源风险："),
    ("权限风险：共享盘或高权限服务账户可能将无权内容写入本地索引。", "权限风险："),
    ("部署风险：企业杀毒、代理、应用白名单和无管理员权限可能阻止后台组件运行。", "部署风险："),
):
    add_bullet(doc, text, bold_lead=lead)

add_page_break(doc)

add_heading(doc, "12. 最终建议", 1)
add_callout(doc, "推荐路线", "先交付一个无需生成式大模型、完全本地运行的高质量混合搜索 MVP；以真实语料评测决定是否增加 Reranker、OCR 和本地问答。", fill=PALE_GREEN, accent=GREEN)

add_heading(doc, "第一阶段：可用的本地搜索", 2)
for item in (
    "完成目录扫描、常见格式解析、结构化分段和增量更新。",
    "使用 SQLite FTS5 + bge-m3/ONNX + HNSW 实现混合检索。",
    "提供可追溯结果、索引状态、失败列表和 Windows 打开操作。",
    "完成性能、安全和至少 100 条业务查询的相关性验收。",
):
    add_bullet(doc, item)

add_heading(doc, "第二阶段：按证据增强", 2)
for item in (
    "若 Top 10 排序仍不足，再引入本地 Reranker。",
    "若扫描件占比较高，再增加 OCR 和独立资源调度。",
    "若用户确有综合答案需求，再加入引用来源的本地 RAG 问答。",
    "若要多人共享或达到百万级以上文档，再评估 Qdrant/OpenSearch 和服务端架构。",
):
    add_bullet(doc, item)

add_heading(doc, "建议技术基线", 2)
add_table(
    doc,
    ["层次", "推荐基线"],
    [
        ("Windows 客户端", ".NET 8 + WPF + CommunityToolkit.Mvvm"),
        ("本地服务", ".NET Worker；解析生态需要时引入打包 Python 运行时"),
        ("语义模型", "bge-m3，ONNX Runtime CPU 默认，GPU 可选"),
        ("索引", "SQLite 元数据 + FTS5 关键词 + HNSW 向量"),
        ("排序", "RRF 融合；Reranker 作为评测后的增强项"),
        ("安全", "无默认外传、仅本机监听、只读原文件、索引可清理"),
    ],
    [2300, 7060],
    header_fill=LIGHT_BLUE,
)

add_body(doc, "这条路线能够在控制部署复杂度和数据风险的同时，优先验证真正决定项目价值的三件事：文档是否解析完整、搜索结果是否相关、用户是否能快速回到原始证据。", bold_lead="这条路线")

# Document properties
doc.core_properties.title = "Windows 本地文档智能搜索 - 实现方案与技术建议"
doc.core_properties.subject = "Windows 本地文档智能搜索供应商评估与实施建议"
doc.core_properties.author = ""
doc.core_properties.keywords = "Windows, 本地搜索, 语义搜索, Embedding, 混合检索, MVP"
doc.core_properties.comments = "基于 2026-08-13 供应商评估版需求文档整理"

doc.save(OUT)
print(OUT)
