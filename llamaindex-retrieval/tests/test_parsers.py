from pathlib import Path

import fitz
from docx import Document as DocxDocument

from llamaindex_retrieval.parsers import parse_uploaded_file


def test_parse_markdown_preserves_management_metadata(tmp_path: Path) -> None:
    path = tmp_path / "policy.md"
    path.write_text("# 报销制度\n\n提交发票。", encoding="utf-8")

    documents = parse_uploaded_file(path, "file-1", "kb-1")

    assert len(documents) == 1
    assert documents[0].metadata["title"] == "报销制度"
    assert documents[0].metadata["file_id"] == "file-1"
    assert documents[0].metadata["knowledge_base_id"] == "kb-1"


def test_parse_markdown_creates_one_document_per_answered_question(tmp_path: Path) -> None:
    path = tmp_path / "questions.md"
    path.write_text(
        """# 问答汇编

#### 社区影响力

**问题 ID：** 282

**问题：**

RWKV 社区影响力如何？

**答案：**

GitHub 有超过 600 个相关项目，Discord 有 9600 多名成员。

---

#### 尚未回答

**问题 ID：** 283

**问题：**

这个问题有答案吗？

**答案状态：** 暂无答案

---
""",
        encoding="utf-8",
    )

    documents = parse_uploaded_file(path, "file-qa", "kb-qa")

    assert len(documents) == 1
    document = documents[0]
    assert document.metadata["title"] == "社区影响力"
    assert document.metadata["question_id"] == "282"
    assert document.metadata["kind"] == "qa-markdown"
    assert document.metadata["answer_status"] == "answered"
    assert document.metadata["full_answer"].startswith("GitHub")
    assert document.metadata["document_id"] != "file-qa"
    assert "尚未回答" not in document.text


def test_parse_markdown_preserves_multiple_answers(tmp_path: Path) -> None:
    path = tmp_path / "multiple.md"
    path.write_text(
        """#### 一个问题

**问题：**

问题正文

**答案 1：**

第一个答案

**答案 2：**

第二个答案

---
""",
        encoding="utf-8",
    )

    documents = parse_uploaded_file(path, "file-multiple", "kb-1")

    assert len(documents) == 1
    assert documents[0].metadata["full_answer"] == (
        "答案 1：\n第一个答案\n\n答案 2：\n第二个答案"
    )


def test_parse_pdf_creates_one_document_per_text_page(tmp_path: Path) -> None:
    path = tmp_path / "guide.pdf"
    pdf = fitz.open()
    first_page = pdf.new_page()
    first_page.insert_text((72, 72), "First page")
    second_page = pdf.new_page()
    second_page.insert_text((72, 72), "Second page")
    pdf.save(path)
    pdf.close()

    documents = parse_uploaded_file(path, "file-2", "kb-2")

    assert len(documents) == 2
    assert documents[0].metadata["page"] == 1
    assert documents[1].metadata["page"] == 2


def test_parse_docx_extracts_paragraphs_and_tables(tmp_path: Path) -> None:
    path = tmp_path / "manual.docx"
    word = DocxDocument()
    word.add_heading("员工手册", level=1)
    word.add_paragraph("请遵守公司制度。")
    table = word.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "项目"
    table.cell(0, 1).text = "要求"
    word.save(path)

    documents = parse_uploaded_file(path, "file-3", "kb-3")

    assert len(documents) == 1
    assert documents[0].metadata["title"] == "员工手册"
    assert "请遵守公司制度" in documents[0].text
    assert "项目\t要求" in documents[0].text
