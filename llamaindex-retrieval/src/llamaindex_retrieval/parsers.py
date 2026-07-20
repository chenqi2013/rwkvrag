import hashlib
import re
from pathlib import Path

import fitz
from docx import Document as DocxDocument
from llama_index.core import Document

SUPPORTED_EXTENSIONS = {".md", ".markdown", ".mdx", ".pdf", ".docx"}
QA_BLOCK_PATTERN = re.compile(
    r"^####[ \t]+(?P<title>.+?)\s*$\n(?P<body>.*?)(?=^---[ \t]*$)",
    re.MULTILINE | re.DOTALL,
)
QA_QUESTION_PATTERN = re.compile(
    r"^\*\*问题：\*\*[ \t]*$\n(?P<content>.*?)(?=^\*\*(?:标签|答案(?: \d+)?|答案状态)：\*\*|\Z)",
    re.MULTILINE | re.DOTALL,
)
QA_ANSWER_PATTERN = re.compile(
    r"^\*\*答案(?: (?P<number>\d+))?：\*\*[ \t]*$",
    re.MULTILINE,
)
QA_ID_PATTERN = re.compile(r"^\*\*问题 ID：\*\*[ \t]*(?P<id>[^\n]+)$", re.MULTILINE)


class UnsupportedDocumentError(ValueError):
    pass


class EmptyDocumentError(ValueError):
    pass


def parsed_document_id(file_id: str, part: str) -> str:
    value = f"{file_id}\0{part}".encode("utf-8")
    return hashlib.sha256(value).hexdigest()[:24]


def markdown_title(content: str, fallback: str) -> str:
    for line in content.splitlines():
        stripped = line.strip()
        if stripped.startswith("# "):
            return stripped.removeprefix("# ").strip() or fallback
    return fallback


def document_metadata(
    *,
    file_id: str,
    knowledge_base_id: str,
    filename: str,
    title: str,
    uri: str,
    kind: str,
    document_id: str,
    page: int | None = None,
) -> dict[str, str | int]:
    metadata: dict[str, str | int] = {
        "document_id": document_id,
        "file_id": file_id,
        "knowledge_base_id": knowledge_base_id,
        "source": "uploaded-document",
        "title": title,
        "uri": uri,
        "file": filename,
        "kind": kind,
    }
    if page is not None:
        metadata["page"] = page
    return metadata


def llama_document(
    *,
    text: str,
    metadata: dict[str, str | int],
) -> Document:
    excluded = [
        "document_id",
        "file_id",
        "knowledge_base_id",
        "source",
        "uri",
        "file",
        "kind",
        "question_id",
        "question",
        "full_answer",
        "answer_status",
    ]
    return Document(
        id_=str(metadata["document_id"]),
        text=text,
        metadata=metadata,
        excluded_embed_metadata_keys=excluded,
        excluded_llm_metadata_keys=excluded,
    )


def parse_qa_markdown(
    content: str,
    path: Path,
    file_id: str,
    knowledge_base_id: str,
) -> tuple[int, list[Document]]:
    documents: list[Document] = []
    matched_blocks = 0
    for block_index, match in enumerate(QA_BLOCK_PATTERN.finditer(content), start=1):
        body = match.group("body").strip()
        question_match = QA_QUESTION_PATTERN.search(body)
        if question_match is None:
            continue
        matched_blocks += 1
        answers = []
        answer_matches = list(QA_ANSWER_PATTERN.finditer(body))
        for answer_index, answer_match in enumerate(answer_matches):
            answer_end = (
                answer_matches[answer_index + 1].start()
                if answer_index + 1 < len(answer_matches)
                else len(body)
            )
            answer = body[answer_match.end() : answer_end].strip()
            if answer:
                answers.append(answer)
        if not answers:
            continue

        title = match.group("title").strip()
        question = question_match.group("content").strip()
        question_id_match = QA_ID_PATTERN.search(body)
        question_id = (
            question_id_match.group("id").strip()
            if question_id_match is not None
            else str(block_index)
        )
        if len(answers) == 1:
            full_answer = answers[0]
        else:
            full_answer = "\n\n".join(
                f"答案 {answer_index}：\n{answer}"
                for answer_index, answer in enumerate(answers, start=1)
            )
        document_id = parsed_document_id(file_id, f"qa:{question_id}")
        metadata = document_metadata(
            file_id=file_id,
            knowledge_base_id=knowledge_base_id,
            filename=path.name,
            title=title,
            uri=f"{path.resolve().as_uri()}#question-{question_id}",
            kind="qa-markdown",
            document_id=document_id,
        )
        metadata.update(
            {
                "question_id": question_id,
                "question": question,
                "full_answer": full_answer,
                "answer_status": "answered",
            }
        )
        text = f"问题标题：{title}\n\n问题：{question}\n\n答案：\n{full_answer}"
        documents.append(llama_document(text=text, metadata=metadata))
    return matched_blocks, documents


def parse_markdown(
    path: Path,
    file_id: str,
    knowledge_base_id: str,
) -> list[Document]:
    content = path.read_text(encoding="utf-8").strip()
    if not content:
        raise EmptyDocumentError("Markdown 文件没有可索引内容")
    qa_blocks, qa_documents = parse_qa_markdown(
        content,
        path,
        file_id,
        knowledge_base_id,
    )
    if qa_blocks:
        if not qa_documents:
            raise EmptyDocumentError("结构化问答 Markdown 没有已回答问题")
        return qa_documents
    title = markdown_title(content, path.stem)
    document_id = parsed_document_id(file_id, "document")
    metadata = document_metadata(
        file_id=file_id,
        knowledge_base_id=knowledge_base_id,
        filename=path.name,
        title=title,
        uri=path.resolve().as_uri(),
        kind="markdown",
        document_id=document_id,
    )
    return [llama_document(text=content, metadata=metadata)]


def parse_pdf(
    path: Path,
    file_id: str,
    knowledge_base_id: str,
) -> list[Document]:
    documents: list[Document] = []
    with fitz.open(path) as pdf:
        for page_index, page in enumerate(pdf):
            content = page.get_text("text", sort=True).strip()
            if not content:
                continue
            page_number = page_index + 1
            document_id = parsed_document_id(file_id, f"page:{page_number}")
            metadata = document_metadata(
                file_id=file_id,
                knowledge_base_id=knowledge_base_id,
                filename=path.name,
                title=path.stem,
                uri=f"{path.resolve().as_uri()}#page={page_number}",
                kind="pdf",
                document_id=document_id,
                page=page_number,
            )
            documents.append(llama_document(text=content, metadata=metadata))
    if not documents:
        raise EmptyDocumentError("PDF 没有可提取文字；扫描版 PDF 需要先执行 OCR")
    return documents


def parse_docx(
    path: Path,
    file_id: str,
    knowledge_base_id: str,
) -> list[Document]:
    word = DocxDocument(path)
    paragraphs = [paragraph.text.strip() for paragraph in word.paragraphs if paragraph.text.strip()]
    tables: list[str] = []
    for table in word.tables:
        rows = [
            "\t".join(cell.text.strip() for cell in row.cells)
            for row in table.rows
            if any(cell.text.strip() for cell in row.cells)
        ]
        if rows:
            tables.append("\n".join(rows))
    content = "\n\n".join([*paragraphs, *tables]).strip()
    if not content:
        raise EmptyDocumentError("DOCX 文件没有可索引内容")
    title = next(
        (
            paragraph.text.strip()
            for paragraph in word.paragraphs
            if paragraph.style.name.startswith("Heading") and paragraph.text.strip()
        ),
        path.stem,
    )
    document_id = parsed_document_id(file_id, "document")
    metadata = document_metadata(
        file_id=file_id,
        knowledge_base_id=knowledge_base_id,
        filename=path.name,
        title=title,
        uri=path.resolve().as_uri(),
        kind="docx",
        document_id=document_id,
    )
    return [llama_document(text=content, metadata=metadata)]


def parse_uploaded_file(
    path: Path,
    file_id: str,
    knowledge_base_id: str,
) -> list[Document]:
    extension = path.suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        raise UnsupportedDocumentError(f"不支持的文件类型：{extension or '无扩展名'}")
    if extension in {".md", ".markdown", ".mdx"}:
        return parse_markdown(path, file_id, knowledge_base_id)
    if extension == ".pdf":
        return parse_pdf(path, file_id, knowledge_base_id)
    return parse_docx(path, file_id, knowledge_base_id)
